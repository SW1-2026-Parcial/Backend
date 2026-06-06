"""
Reportes dinámicos — Ciclo 2.
El gerente/admin escribe un prompt en lenguaje natural, la IA parsea los criterios,
y el backend genera un Excel o Word que se descarga directamente (sin almacenar en Azure).
Incluye endpoint /chat para conversación con memoria de sesión.
"""
import io
import json
import logging
import uuid
from datetime import datetime, timezone
from typing import Any, List, Optional

from beanie import PydanticObjectId
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from core.security import get_current_user, require_admin_or_supervisor
from core.exceptions import BusinessException
from models.user import User
from models.tramite import Tramite, EstadoTramite
from models.tramite_event import TramiteEvent
from models.task import Task
from models.departamento import Departamento
from models.politica import Politica
from modules.intelligence.llm.client import chat_completion_json

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/reportes", tags=["reportes"])


# ── Schemas ───────────────────────────────────────────────────────────────────

class CriteriosReporte(BaseModel):
    """Criterios estructurados parseados del prompt por la IA."""
    titulo: str = "Reporte BPM"
    formato: str = "EXCEL"                     # EXCEL | WORD
    estado: Optional[str] = None               # ACTIVE | COMPLETED | REJECTED | CANCELLED
    departamentoId: Optional[str] = None
    politicaId: Optional[str] = None
    fechaDesde: Optional[str] = None           # ISO date string
    fechaHasta: Optional[str] = None
    columnas: List[str] = []                   # campos a incluir
    ordenarPor: str = "startedAt"


class ParsearPromptRequest(BaseModel):
    prompt: str


class ParsearPromptResponse(BaseModel):
    criterios: CriteriosReporte
    valido: bool
    advertencias: List[str] = []


class GenerarReporteRequest(BaseModel):
    criterios: CriteriosReporte


# ── Helpers de datos ──────────────────────────────────────────────────────────

async def _consultar_tramites(criterios: CriteriosReporte) -> List[dict]:
    """Consulta trámites en MongoDB según los criterios."""
    query: dict = {}

    if criterios.estado:
        try:
            query["status"] = EstadoTramite(criterios.estado)
        except ValueError:
            pass

    if criterios.politicaId:
        query["politicaId"] = criterios.politicaId

    if criterios.fechaDesde or criterios.fechaHasta:
        rango: dict = {}
        if criterios.fechaDesde:
            rango["$gte"] = datetime.fromisoformat(criterios.fechaDesde)
        if criterios.fechaHasta:
            rango["$lte"] = datetime.fromisoformat(criterios.fechaHasta)
        query["startedAt"] = rango

    tramites = await Tramite.find(query).sort("-startedAt").to_list()

    # Enriquecer con nombre de política y departamento si aplica
    politicas_cache: dict = {}
    departamentos_cache: dict = {}

    rows = []
    for t in tramites:
        # Cache de políticas
        pol_nombre = politicas_cache.get(t.politicaId)
        if pol_nombre is None:
            pol = await Politica.get(PydanticObjectId(t.politicaId))
            pol_nombre = pol.nombre if pol else t.politicaId
            politicas_cache[t.politicaId] = pol_nombre

        duracion_horas: Optional[float] = None
        if t.startedAt and t.completedAt:
            delta = t.completedAt - t.startedAt
            duracion_horas = round(delta.total_seconds() / 3600, 1)

        row = {
            "ticket":         t.ticketNumber or "",
            "politica":       pol_nombre,
            "estado":         t.status.value,
            "prioridad":      t.prioridad.value,
            "iniciado_por":   t.initiatedBy or "",
            "fecha_inicio":   t.startedAt.strftime("%Y-%m-%d %H:%M") if t.startedAt else "",
            "fecha_fin":      t.completedAt.strftime("%Y-%m-%d %H:%M") if t.completedAt else "",
            "duracion_horas": duracion_horas or "",
        }
        rows.append(row)

    return rows


# ── Generadores ───────────────────────────────────────────────────────────────

def _generar_excel(titulo: str, rows: List[dict]) -> bytes:
    """Genera un archivo Excel con openpyxl."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Reporte"

    # Título
    ws.merge_cells("A1:H1")
    titulo_cell = ws["A1"]
    titulo_cell.value = titulo
    titulo_cell.font = Font(bold=True, size=14)
    titulo_cell.alignment = Alignment(horizontal="center")

    ws.append([])  # fila vacía

    if not rows:
        ws.append(["Sin datos para los criterios seleccionados"])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    # Encabezados
    headers = list(rows[0].keys())
    header_labels = {
        "ticket": "Ticket", "politica": "Política", "estado": "Estado",
        "prioridad": "Prioridad", "iniciado_por": "Iniciado Por",
        "fecha_inicio": "Fecha Inicio", "fecha_fin": "Fecha Fin",
        "duracion_horas": "Duración (hs)",
    }
    header_row = [header_labels.get(h, h.title()) for h in headers]

    ws.append(header_row)
    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for col_idx, _ in enumerate(header_row, start=1):
        cell = ws.cell(row=ws.max_row, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # Datos con color alternado
    fill_alt = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")
    for i, row in enumerate(rows):
        ws.append([row.get(h, "") for h in headers])
        if i % 2 == 0:
            for col_idx in range(1, len(headers) + 1):
                ws.cell(row=ws.max_row, column=col_idx).fill = fill_alt

    # Auto-ancho de columnas
    for col_idx, _ in enumerate(headers, start=1):
        col_letter = get_column_letter(col_idx)
        max_len = max(
            (len(str(ws.cell(r, col_idx).value or "")) for r in range(3, ws.max_row + 1)),
            default=10,
        )
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _generar_word(titulo: str, rows: List[dict]) -> bytes:
    """Genera un archivo Word con python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Título
    titulo_p = doc.add_heading(titulo, level=1)
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')} | "
        f"Total registros: {len(rows)}"
    )
    doc.add_paragraph("")

    if not rows:
        doc.add_paragraph("Sin datos para los criterios seleccionados.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    headers = list(rows[0].keys())
    header_labels = {
        "ticket": "Ticket", "politica": "Política", "estado": "Estado",
        "prioridad": "Prioridad", "iniciado_por": "Iniciado Por",
        "fecha_inicio": "Fecha Inicio", "fecha_fin": "Fecha Fin",
        "duracion_horas": "Duración (hs)",
    }

    # Tabla
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = header_labels.get(h, h.title())
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Datos
    for row in rows:
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row.get(h, ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/parse-prompt", response_model=ParsearPromptResponse)
async def parsear_prompt(
    body: ParsearPromptRequest,
    _: User = Depends(require_admin_or_supervisor),
):
    """
    Usa el LLM (OpenRouter) para parsear un prompt en lenguaje natural
    a criterios estructurados para generación de reportes.
    Si el LLM falla, cae al parser básico por palabras clave.
    """
    system_prompt = """Eres un parser de reportes BPM. El usuario describe en lenguaje natural qué reporte quiere.
Tu trabajo es extraer los criterios estructurados y devolver SOLO un JSON con esta estructura exacta:
{
  "criterios": {
    "titulo": "título descriptivo del reporte",
    "formato": "EXCEL o WORD",
    "estado": "ACTIVE | COMPLETED | REJECTED | CANCELLED | null",
    "departamentoId": "id del departamento o null",
    "politicaId": "id de la política o null",
    "fechaDesde": "YYYY-MM-DD o null",
    "fechaHasta": "YYYY-MM-DD o null",
    "columnas": [],
    "ordenarPor": "startedAt"
  },
  "valido": true,
  "advertencias": ["advertencia si algo no se pudo interpretar"]
}

Reglas:
- Si el usuario dice "excel" usa EXCEL, si dice "word/doc" usa WORD. Por defecto EXCEL.
- Los estados posibles son: ACTIVE, COMPLETED, REJECTED, CANCELLED.
- Si el usuario menciona fechas relativas como "este mes" o "mayo", calcúlalas (hoy es """ + datetime.now().strftime('%Y-%m-%d') + """).
- Si algo es ambiguo, pon null y añade una advertencia.
- SIEMPRE devuelve JSON válido, nada más."""

    try:
        data = await chat_completion_json(
            system_prompt=system_prompt,
            user_message=body.prompt,
            temperature=0.1,
        )
        criterios_data = data.get("criterios", data)
        return ParsearPromptResponse(
            criterios=CriteriosReporte(**criterios_data),
            valido=data.get("valido", True),
            advertencias=data.get("advertencias", []),
        )
    except Exception as e:
        logger.warning("LLM no disponible para parse-prompt: %s — usando parser básico", e)
        criterios = _parsear_basico(body.prompt)
        return ParsearPromptResponse(
            criterios=criterios,
            valido=True,
            advertencias=["Análisis con parser básico — LLM no disponible"],
        )


def _parsear_basico(prompt: str) -> CriteriosReporte:
    """Parser de fallback por palabras clave cuando sp1-ai no responde."""
    p = prompt.lower()
    criterios = CriteriosReporte(titulo=f"Reporte: {prompt[:60]}")

    if "excel" in p:
        criterios.formato = "EXCEL"
    elif "word" in p or "doc" in p:
        criterios.formato = "WORD"

    if "completad" in p:
        criterios.estado = "COMPLETED"
    elif "activ" in p or "pendient" in p:
        criterios.estado = "ACTIVE"
    elif "rechazad" in p or "cancelad" in p:
        criterios.estado = "REJECTED"

    return criterios


@router.post("/generar")
async def generar_reporte(
    body: GenerarReporteRequest,
    _: User = Depends(require_admin_or_supervisor),
):
    """
    Genera el reporte (Excel o Word) y lo retorna como descarga directa.
    No almacena nada en Azure — el usuario descarga el archivo en el momento.
    """
    criterios = body.criterios
    rows = await _consultar_tramites(criterios)

    titulo = criterios.titulo or "Reporte BPM"
    ahora = datetime.now().strftime("%Y%m%d_%H%M")

    if criterios.formato == "WORD":
        contenido = _generar_word(titulo, rows)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"reporte_{ahora}.docx"
    else:
        contenido = _generar_excel(titulo, rows)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename = f"reporte_{ahora}.xlsx"

    logger.info("Reporte generado: %s (%d filas)", filename, len(rows))

    return StreamingResponse(
        io.BytesIO(contenido),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Chat conversacional para reportes ────────────────────────────────────────

_reporte_sessions: dict[str, list[dict]] = {}
MAX_REPORTE_HISTORY = 20

REPORTE_SYSTEM_PROMPT = """Eres un asistente de reportes del Sistema BPM SP1. Tu trabajo es ayudar al usuario a generar reportes de trámites conversando en lenguaje natural.

FLUJO:
1. El usuario describe qué reporte quiere.
2. Vos analizás si tenés toda la información necesaria.
3. Si falta algo, preguntás de forma amable y concisa.
4. Cuando tenés todo, devolvés los criterios confirmados.

CRITERIOS QUE NECESITÁS COMPLETAR:
- formato: EXCEL o WORD (por defecto EXCEL)
- estado: ACTIVE | COMPLETED | REJECTED | CANCELLED | null (todos)
- fechaDesde / fechaHasta: YYYY-MM-DD o null
- politicaId: id de la política o null (todas)
- departamentoId: id del departamento o null (todos)
- titulo: título descriptivo del reporte

REGLAS:
- Respondé siempre en español, de forma concisa y amable.
- Usá **negritas** para resaltar datos importantes.
- Si el usuario dice fechas relativas como "este mes", "mayo", "la semana pasada", calculá las fechas exactas (hoy es {today}).
- Si te falta información crítica (al menos estado O período), preguntá.
- Si no se especifica formato, usá EXCEL por defecto.
- Cuando tengas suficiente información, confirmá todo y marcá listo.

FORMATO DE RESPUESTA — SIEMPRE JSON (usa llaves simples, NO dobles):
{
  "mensaje": "texto de respuesta al usuario",
  "criterios": {
    "titulo": "...",
    "formato": "EXCEL",
    "estado": null,
    "departamentoId": null,
    "politicaId": null,
    "fechaDesde": null,
    "fechaHasta": null,
    "columnas": [],
    "ordenarPor": "startedAt"
  },
  "listo": false,
  "camposFaltantes": ["estado", "fechas"]
}

Cuando tengas todo y el usuario confirme: pon "listo": true y "camposFaltantes": [].
"""


class ReporteChatRequest(BaseModel):
    mensaje: str = ""          # vacío permitido — /chat/clear no necesita mensaje
    sessionId: Optional[str] = None


class ReporteChatResponse(BaseModel):
    mensaje: str
    criterios: Optional[CriteriosReporte] = None
    listo: bool = False
    camposFaltantes: list[str] = []
    sessionId: str


@router.post("/chat", response_model=ReporteChatResponse)
async def reporte_chat(
    body: ReporteChatRequest,
    _: User = Depends(require_admin_or_supervisor),
):
    session_id = body.sessionId or str(uuid.uuid4())

    if session_id not in _reporte_sessions:
        _reporte_sessions[session_id] = []
    history = _reporte_sessions[session_id]

    history.append({"role": "user", "content": body.mensaje})
    if len(history) > MAX_REPORTE_HISTORY:
        _reporte_sessions[session_id] = history[-MAX_REPORTE_HISTORY:]
        history = _reporte_sessions[session_id]

    system = REPORTE_SYSTEM_PROMPT.replace("{today}", datetime.now().strftime("%Y-%m-%d"))

    try:
        result = await chat_completion_json(
            system_prompt=system,
            user_message=body.mensaje,
            history=history[:-1],
            temperature=0.3,
        )

        msg = result.get("mensaje", "No pude procesar tu solicitud.")
        history.append({"role": "assistant", "content": msg})

        criterios_data = result.get("criterios")
        criterios = CriteriosReporte(**criterios_data) if criterios_data else None

        return ReporteChatResponse(
            mensaje=msg,
            criterios=criterios,
            listo=result.get("listo", False),
            camposFaltantes=result.get("camposFaltantes", []),
            sessionId=session_id,
        )
    except Exception as e:
        logger.error("Error en reporte chat: %s", e)
        fallback = "Hubo un error procesando tu mensaje. ¿Podés reformularlo?"
        history.append({"role": "assistant", "content": fallback})
        return ReporteChatResponse(
            mensaje=fallback,
            sessionId=session_id,
        )


@router.post("/chat/clear")
async def reporte_chat_clear(
    body: ReporteChatRequest,
    _: User = Depends(require_admin_or_supervisor),
):
    if body.sessionId:
        _reporte_sessions.pop(body.sessionId, None)
    return {"ok": True}
