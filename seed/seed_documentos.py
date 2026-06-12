"""
Seed de documentos — sube archivos a Azure Blob Storage y crea registros en MongoDB.

Genera documentos realistas (PDF, DOCX, XLSX) para un subconjunto de trámites existentes.
Cada documento queda correctamente linkeado: politicaId, versionPoliticaId, tramiteId.

Uso:
    cd sp1-backend-py
    python seed/seed_documentos.py              # 50 trámites random, 1-3 docs c/u
    python seed/seed_documentos.py --count 100  # 100 trámites
    python seed/seed_documentos.py --all        # todos los 300 trámites
"""
import asyncio
import io
import json
import logging
import os
import random
import sys
import uuid
from datetime import datetime, timezone, timedelta
from pathlib import Path

# ── Configurar path para imports del proyecto ────────────────────────────────
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from azure.storage.blob import BlobServiceClient, ContentSettings
from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parent.parent / ".env")

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ── Config Azure ─────────────────────────────────────────────────────────────
CONN_STR = os.getenv("AZURE_STORAGE_CONNECTION_STRING", "")
CONTAINER = os.getenv("AZURE_STORAGE_CONTAINER", "documentos")

# ── Plantillas de documentos por tipo de política ────────────────────────────
# politicaId → lista de posibles documentos
DOCS_POR_POLITICA = {
    "cc1000000000000000000001": [  # reporte_incidencia_tecnica
        ("reporte_incidencia.pdf", "pdf"),
        ("evidencia_fotografica.jpg", "jpg"),
        ("diagnostico_tecnico.docx", "docx"),
        ("cotizacion_reparacion.xlsx", "xlsx"),
        ("formulario_incidencia.pdf", "pdf"),
    ],
    "cc1000000000000000000002": [  # evaluacion_proveedor
        ("evaluacion_proveedor.pdf", "pdf"),
        ("contrato_proveedor.docx", "docx"),
        ("tabla_comparativa_precios.xlsx", "xlsx"),
        ("certificado_calidad.pdf", "pdf"),
        ("informe_auditoria.docx", "docx"),
    ],
    "cc1000000000000000000003": [  # solicitud_vacaciones
        ("solicitud_vacaciones.pdf", "pdf"),
        ("calendario_equipo.xlsx", "xlsx"),
        ("aprobacion_jefe.pdf", "pdf"),
        ("constancia_dias_disponibles.docx", "docx"),
    ],
    "cc1000000000000000000004": [  # proceso_contratacion
        ("cv_candidato.pdf", "pdf"),
        ("carta_presentacion.docx", "docx"),
        ("evaluacion_tecnica.xlsx", "xlsx"),
        ("contrato_laboral.docx", "docx"),
        ("cedula_identidad.jpg", "jpg"),
        ("certificado_estudios.pdf", "pdf"),
    ],
    "cc1000000000000000000005": [  # reclamo_cliente
        ("formulario_reclamo.pdf", "pdf"),
        ("evidencia_problema.jpg", "jpg"),
        ("factura_compra.pdf", "pdf"),
        ("respuesta_oficial.docx", "docx"),
        ("acta_resolucion.pdf", "pdf"),
    ],
}

# Fallback genérico
DOCS_GENERICOS = [
    ("documento_adjunto.pdf", "pdf"),
    ("nota_interna.docx", "docx"),
    ("datos_registro.xlsx", "xlsx"),
]

MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "jpg": "image/jpeg",
    "png": "image/png",
    "txt": "text/plain",
}

USUARIOS_FUNCIONARIOS = [
    "69f1d907b60ba6bbff7bccc3",
    "69f1d907b60ba6bbff7bccc4",
    "69f1d907b60ba6bbff7bccc5",
    "69f1d907b60ba6bbff7bccc6",
]


# ── Generadores de archivos dummy ────────────────────────────────────────────

def _generar_pdf(nombre: str, tramite_ticket: str) -> bytes:
    """Genera un PDF mínimo válido con texto."""
    # PDF mínimo válido con contenido legible
    contenido = f"Documento: {nombre}\nTramite: {tramite_ticket}\nFecha: {datetime.now().strftime('%d/%m/%Y')}"
    # PDF 1.4 mínimo con texto
    text_stream = f"BT /F1 12 Tf 50 750 Td ({contenido.replace(chr(10), ') Tj T* (')}) Tj ET"
    stream_bytes = text_stream.encode("latin-1")
    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n"
        b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"5 0 obj<</Length " + str(len(stream_bytes)).encode() + b">>stream\n"
        + stream_bytes + b"\nendstream\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000340 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n0\n%%EOF"
    )
    return pdf


def _generar_docx(nombre: str, tramite_ticket: str) -> bytes:
    """Genera un DOCX mínimo válido."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(nombre, level=1)
        doc.add_paragraph(f"Trámite: {tramite_ticket}")
        doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("Este documento fue generado automáticamente como dato de prueba del sistema BPM SP1.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # Fallback: archivo de texto simple
        return f"Documento: {nombre}\nTramite: {tramite_ticket}\n".encode()


def _generar_xlsx(nombre: str, tramite_ticket: str) -> bytes:
    """Genera un XLSX mínimo válido."""
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Datos"
        ws.append(["Campo", "Valor"])
        ws.append(["Documento", nombre])
        ws.append(["Trámite", tramite_ticket])
        ws.append(["Fecha", datetime.now().strftime("%d/%m/%Y")])
        ws.append(["Estado", "Generado automáticamente"])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except ImportError:
        return b"placeholder xlsx content"


def _generar_jpg_placeholder(nombre: str) -> bytes:
    """Genera un JPEG mínimo válido (1x1 pixel gris)."""
    # JFIF mínimo válido — 1x1 pixel
    return bytes([
        0xFF, 0xD8, 0xFF, 0xE0, 0x00, 0x10, 0x4A, 0x46, 0x49, 0x46, 0x00, 0x01,
        0x01, 0x00, 0x00, 0x01, 0x00, 0x01, 0x00, 0x00, 0xFF, 0xDB, 0x00, 0x43,
        0x00, 0x08, 0x06, 0x06, 0x07, 0x06, 0x05, 0x08, 0x07, 0x07, 0x07, 0x09,
        0x09, 0x08, 0x0A, 0x0C, 0x14, 0x0D, 0x0C, 0x0B, 0x0B, 0x0C, 0x19, 0x12,
        0x13, 0x0F, 0x14, 0x1D, 0x1A, 0x1F, 0x1E, 0x1D, 0x1A, 0x1C, 0x1C, 0x20,
        0x24, 0x2E, 0x27, 0x20, 0x22, 0x2C, 0x23, 0x1C, 0x1C, 0x28, 0x37, 0x29,
        0x2C, 0x30, 0x31, 0x34, 0x34, 0x34, 0x1F, 0x27, 0x39, 0x3D, 0x38, 0x32,
        0x3C, 0x2E, 0x33, 0x34, 0x32, 0xFF, 0xC0, 0x00, 0x0B, 0x08, 0x00, 0x01,
        0x00, 0x01, 0x01, 0x01, 0x11, 0x00, 0xFF, 0xC4, 0x00, 0x1F, 0x00, 0x00,
        0x01, 0x05, 0x01, 0x01, 0x01, 0x01, 0x01, 0x01, 0x00, 0x00, 0x00, 0x00,
        0x00, 0x00, 0x00, 0x00, 0x01, 0x02, 0x03, 0x04, 0x05, 0x06, 0x07, 0x08,
        0x09, 0x0A, 0x0B, 0xFF, 0xC4, 0x00, 0xB5, 0x10, 0x00, 0x02, 0x01, 0x03,
        0x03, 0x02, 0x04, 0x03, 0x05, 0x05, 0x04, 0x04, 0x00, 0x00, 0x01, 0x7D,
        0x01, 0x02, 0x03, 0x00, 0x04, 0x11, 0x05, 0x12, 0x21, 0x31, 0x41, 0x06,
        0x13, 0x51, 0x61, 0x07, 0x22, 0x71, 0x14, 0x32, 0x81, 0x91, 0xA1, 0x08,
        0x23, 0x42, 0xB1, 0xC1, 0x15, 0x52, 0xD1, 0xF0, 0x24, 0x33, 0x62, 0x72,
        0x82, 0x09, 0x0A, 0x16, 0x17, 0x18, 0x19, 0x1A, 0x25, 0x26, 0x27, 0x28,
        0x29, 0x2A, 0x34, 0x35, 0x36, 0x37, 0x38, 0x39, 0x3A, 0x43, 0x44, 0x45,
        0x46, 0x47, 0x48, 0x49, 0x4A, 0x53, 0x54, 0x55, 0x56, 0x57, 0x58, 0x59,
        0x5A, 0x63, 0x64, 0x65, 0x66, 0x67, 0x68, 0x69, 0x6A, 0x73, 0x74, 0x75,
        0x76, 0x77, 0x78, 0x79, 0x7A, 0x83, 0x84, 0x85, 0x86, 0x87, 0x88, 0x89,
        0x8A, 0x92, 0x93, 0x94, 0x95, 0x96, 0x97, 0x98, 0x99, 0x9A, 0xA2, 0xA3,
        0xA4, 0xA5, 0xA6, 0xA7, 0xA8, 0xA9, 0xAA, 0xB2, 0xB3, 0xB4, 0xB5, 0xB6,
        0xB7, 0xB8, 0xB9, 0xBA, 0xC2, 0xC3, 0xC4, 0xC5, 0xC6, 0xC7, 0xC8, 0xC9,
        0xCA, 0xD2, 0xD3, 0xD4, 0xD5, 0xD6, 0xD7, 0xD8, 0xD9, 0xDA, 0xE1, 0xE2,
        0xE3, 0xE4, 0xE5, 0xE6, 0xE7, 0xE8, 0xE9, 0xEA, 0xF1, 0xF2, 0xF3, 0xF4,
        0xF5, 0xF6, 0xF7, 0xF8, 0xF9, 0xFA, 0xFF, 0xDA, 0x00, 0x08, 0x01, 0x01,
        0x00, 0x00, 0x3F, 0x00, 0x7B, 0x94, 0x11, 0x00, 0x00, 0x00, 0x00, 0x00,
        0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0x00, 0xFF, 0xD9,
    ])


def generar_contenido(nombre: str, ext: str, ticket: str) -> bytes:
    """Genera contenido de archivo según extensión."""
    if ext == "pdf":
        return _generar_pdf(nombre, ticket)
    elif ext == "docx":
        return _generar_docx(nombre, ticket)
    elif ext == "xlsx":
        return _generar_xlsx(nombre, ticket)
    elif ext in ("jpg", "jpeg", "png"):
        return _generar_jpg_placeholder(nombre)
    else:
        return f"Contenido de {nombre} para {ticket}".encode()


# ── Upload a Azure ───────────────────────────────────────────────────────────

def upload_to_azure(blob_name: str, contenido: bytes, mime_type: str) -> int:
    """Sube un blob a Azure Storage. Retorna tamaño en bytes."""
    client = BlobServiceClient.from_connection_string(CONN_STR)
    blob_client = client.get_blob_client(container=CONTAINER, blob=blob_name)
    blob_client.upload_blob(
        contenido,
        overwrite=True,
        content_settings=ContentSettings(content_type=mime_type),
    )
    return len(contenido)


# ── Main ─────────────────────────────────────────────────────────────────────

async def main():
    import argparse
    parser = argparse.ArgumentParser(description="Seed documentos en Azure + MongoDB")
    parser.add_argument("--count", type=int, default=50, help="Cuántos trámites poblar con docs (default 50)")
    parser.add_argument("--all", action="store_true", help="Poblar todos los trámites")
    parser.add_argument("--dry-run", action="store_true", help="Solo mostrar qué haría, sin subir")
    args = parser.parse_args()

    # Cargar trámites del seed
    seed_dir = Path(__file__).resolve().parent
    with open(seed_dir / "tramites.json") as f:
        tramites = json.load(f)

    if args.all:
        seleccionados = tramites
    else:
        seleccionados = random.sample(tramites, min(args.count, len(tramites)))

    logger.info("Generando documentos para %d trámites...", len(seleccionados))

    # Preparar documentos a crear
    docs_to_insert = []
    total_uploaded = 0

    for tramite in seleccionados:
        pol_id = tramite["politicaId"]
        ver_id = tramite["versionPoliticaId"]
        tram_id = tramite["_id"]["$oid"]
        ticket = tramite["ticketNumber"]
        started_at = tramite.get("startedAt", {}).get("$date", "2025-06-01T00:00:00Z")

        # Elegir docs para esta política
        plantillas = DOCS_POR_POLITICA.get(pol_id, DOCS_GENERICOS)
        num_docs = random.randint(1, min(3, len(plantillas)))
        docs_elegidos = random.sample(plantillas, num_docs)

        for nombre_base, ext in docs_elegidos:
            # Nombre con variación: agregar ticket para que sea único
            nombre = f"{nombre_base.rsplit('.', 1)[0]}_{ticket}.{ext}"
            blob_name = f"{pol_id}/{tram_id}/{uuid.uuid4()}.{ext}"
            mime = MIME_TYPES.get(ext, "application/octet-stream")

            if args.dry_run:
                logger.info("  [DRY] %s → %s", nombre, blob_name)
                continue

            # Generar y subir
            contenido = generar_contenido(nombre, ext, ticket)
            tamano = upload_to_azure(blob_name, contenido, mime)
            total_uploaded += 1

            # Fecha del doc: entre startedAt y ahora
            try:
                base_date = datetime.fromisoformat(started_at.replace("Z", "+00:00"))
            except Exception:
                base_date = datetime(2025, 6, 1, tzinfo=timezone.utc)
            doc_date = base_date + timedelta(hours=random.randint(1, 72))

            doc_record = {
                "nombre": nombre,
                "extension": ext,
                "blobName": blob_name,
                "tamano": tamano,
                "mimeType": mime,
                "politicaId": pol_id,
                "versionPoliticaId": ver_id,
                "tramiteId": tram_id,
                "clienteId": tramite.get("initiatedBy"),
                "subidoPorId": random.choice(USUARIOS_FUNCIONARIOS),
                "modificadoPorId": None,
                "permisos": [],
                "activo": True,
                "creadoEn": doc_date.isoformat(),
                "actualizadoEn": None,
            }
            docs_to_insert.append(doc_record)

            if total_uploaded % 25 == 0:
                logger.info("  ... %d archivos subidos a Azure", total_uploaded)

    if args.dry_run:
        logger.info("Dry run completo. %d docs se crearían.", len(docs_to_insert))
        return

    logger.info("Subidos %d archivos a Azure Blob Storage.", total_uploaded)

    # Insertar en MongoDB
    logger.info("Insertando %d registros en MongoDB...", len(docs_to_insert))

    from motor.motor_asyncio import AsyncIOMotorClient

    mongo_uri = os.getenv("MONGODB_URI", "mongodb://localhost:27017")
    db_name = os.getenv("MONGODB_DB_NAME", "swp1_db")

    client = AsyncIOMotorClient(mongo_uri)
    db = client[db_name]
    collection = db["documentos"]

    # Convertir fechas a datetime objects para MongoDB
    for doc in docs_to_insert:
        doc["creadoEn"] = datetime.fromisoformat(doc["creadoEn"])

    if docs_to_insert:
        result = await collection.insert_many(docs_to_insert)
        logger.info("Insertados %d documentos en MongoDB.", len(result.inserted_ids))

    client.close()
    logger.info("✓ Seed completo: %d archivos en Azure, %d registros en MongoDB.", total_uploaded, len(docs_to_insert))


if __name__ == "__main__":
    asyncio.run(main())
